"""File exporter service — generates DOCX and PDF from document text."""
import os
import re
import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

logger = logging.getLogger(__name__)

OUTPUTS_DIR = Path(__file__).parent.parent / "outputs"
OUTPUTS_DIR.mkdir(exist_ok=True)


def _parse_sections(text: str) -> list[dict]:
    """Parse generated text into sections with headings and content."""
    lines = text.strip().split("\n")
    sections = []
    current_section = {"heading": "", "content": []}

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_section["content"]:
                current_section["content"].append("")
            continue

        # Detect section headings (numbered like "1.", "1.1", or ALL CAPS lines)
        is_heading = bool(
            re.match(r"^\d+\.?\s+[A-Z]", stripped)
            or (stripped.isupper() and len(stripped) > 3 and len(stripped) < 120)
        )

        if is_heading:
            if current_section["heading"] or current_section["content"]:
                sections.append(current_section)
            current_section = {"heading": stripped, "content": []}
        else:
            current_section["content"].append(stripped)

    if current_section["heading"] or current_section["content"]:
        sections.append(current_section)

    return sections


def generate_docx(text: str, session_id: str, document_type: str) -> str:
    """Generate a DOCX file from document text. Returns file path."""
    filename = f"{document_type}_{session_id[:8]}.docx"
    filepath = OUTPUTS_DIR / filename

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    sections = _parse_sections(text)

    for i, sec in enumerate(sections):
        heading = sec["heading"]
        content_lines = sec["content"]

        if heading:
            # First section heading is the title
            if i == 0 and ("AGREEMENT" in heading.upper() or "NDA" in heading.upper()):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(heading)
                run.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 0, 0)
                doc.add_paragraph()  # spacer
            else:
                p = doc.add_paragraph()
                run = p.add_run(heading)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

        content_text = "\n".join(content_lines).strip()
        if content_text:
            for para_text in content_text.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.space_after = Pt(6)

    doc.save(str(filepath))
    logger.info(f"DOCX generated: {filepath}")
    return str(filepath)


def generate_pdf(text: str, session_id: str, document_type: str) -> str:
    """Generate a PDF file from document text. Returns file path."""
    filename = f"{document_type}_{session_id[:8]}.pdf"
    filepath = OUTPUTS_DIR / filename

    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=A4,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        leftMargin=1.25 * inch,
        rightMargin=1.25 * inch,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Title"],
        fontSize=16,
        spaceAfter=20,
        alignment=1,  # center
        textColor="black",
    )

    heading_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontSize=12,
        spaceBefore=14,
        spaceAfter=6,
        textColor="black",
        fontName="Helvetica-Bold",
    )

    body_style = ParagraphStyle(
        "DocBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=6,
        textColor="#1a1a1a",
    )

    flowables = []
    sections = _parse_sections(text)

    for i, sec in enumerate(sections):
        heading = sec["heading"]
        content_lines = sec["content"]

        if heading:
            # Escape XML special characters for reportlab
            safe_heading = (
                heading.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            if i == 0 and ("AGREEMENT" in heading.upper() or "NDA" in heading.upper()):
                flowables.append(Paragraph(safe_heading, title_style))
                flowables.append(Spacer(1, 12))
            else:
                flowables.append(Paragraph(safe_heading, heading_style))

        content_text = "\n".join(content_lines).strip()
        if content_text:
            for para_text in content_text.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    safe_text = (
                        para_text.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    flowables.append(Paragraph(safe_text, body_style))

    if flowables:
        doc.build(flowables)
    else:
        # If parsing produced nothing, just dump the raw text
        safe_text = (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )
        flowables.append(Paragraph(safe_text, body_style))
        doc.build(flowables)

    logger.info(f"PDF generated: {filepath}")
    return str(filepath)
